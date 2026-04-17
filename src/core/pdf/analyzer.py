"""
Comprehensive document analysis for multiple file types with OCR support.
Supports PDFs, Word docs, text files, images, and more.
No authentication required.
"""

import io
from datetime import datetime
from typing import Any, Dict

import httpx

from src.logging.logger import logger


class DocumentAnalyzer:
    """Analyze documents of multiple types by extracting text content with OCR.

    Supported formats:
    - PDF (.pdf) - with OCR fallback for scanned docs
    - Word (.docx)
    - Text (.txt, .md)
    - Images (.jpg, .jpeg, .png, .gif, .webp) - with OCR
    - Any text-based format
    """

    def __init__(self):
        self.headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
        }
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        self.ocr_available = self._check_ocr_available()

    def _check_ocr_available(self) -> bool:
        """Check if OCR dependencies are available."""
        import importlib.util

        available = (
            importlib.util.find_spec("easyocr") is not None
            and importlib.util.find_spec("PIL") is not None
        )
        if not available:
            logger.warning(
                "OCR not available - install pillow and easyocr for image/scanned PDF support"
            )
        return available

    def _detect_document_type(self, url: str, content_type: str) -> str:
        """Detect document type from URL and content-type."""
        url_lower = url.lower()

        if ".pdf" in url_lower or "application/pdf" in content_type:
            return "pdf"
        elif ".docx" in url_lower or "application/vnd.openxmlformats" in content_type:
            return "docx"
        elif ".doc" in url_lower or "application/msword" in content_type:
            return "doc"
        elif ".txt" in url_lower or "text/plain" in content_type:
            return "txt"
        elif ".md" in url_lower or "text/markdown" in content_type:
            return "markdown"
        elif any(ext in url_lower for ext in [".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"]):
            return "image"
        elif "image/" in content_type:
            return "image"
        else:
            return "unknown"

    async def analyze_document(
        self, url: str, max_pages: int = 10, extract_metadata: bool = True
    ) -> Dict[str, Any]:
        """
        Download and analyze a document of any supported type.

        Args:
            url: URL of the document
            max_pages: Maximum pages to extract (default: 10, for PDFs)
            extract_metadata: Whether to extract document metadata

        Returns:
            Dictionary with document content and metadata
        """
        try:
            logger.info(f"Downloading document from: {url}")
            async with httpx.AsyncClient(headers=self.headers, timeout=60.0) as client:
                response = await client.get(url, follow_redirects=True)
                response.raise_for_status()

                content_type = response.headers.get("content-type", "")
                content_length = len(response.content)

                if content_length > self.max_file_size:
                    return {
                        "status": "error",
                        "error": f"Document too large: {content_length / 1024 / 1024:.1f}MB (max 50MB)",
                    }

                doc_type = self._detect_document_type(url, content_type)
                logger.info(f"Document type detected: {doc_type}")

                doc_bytes = response.content

            if doc_type == "pdf":
                return await self._analyze_pdf(doc_bytes, url, max_pages, extract_metadata)
            elif doc_type == "docx":
                return await self._analyze_docx(doc_bytes, url, extract_metadata)
            elif doc_type == "image":
                return await self._analyze_image(doc_bytes, url)
            elif doc_type in ["txt", "markdown"]:
                return await self._analyze_text(doc_bytes, url, doc_type)
            else:
                return await self._analyze_text(doc_bytes, url, "unknown")

        except Exception as e:
            logger.error(f"Document analysis failed: {e}")
            return {"status": "error", "error": str(e), "url": url}

    async def _analyze_pdf(
        self, doc_bytes: bytes, url: str, max_pages: int, extract_metadata: bool
    ) -> Dict[str, Any]:
        """Analyze PDF document."""
        try:
            from pypdf import PdfReader
        except ImportError:
            return {"status": "error", "error": "pypdf library not available", "url": url}

        pdf_file = io.BytesIO(doc_bytes)
        reader = PdfReader(pdf_file)

        metadata = {}
        if extract_metadata and reader.metadata:
            metadata = {
                "title": reader.metadata.get("/Title", ""),
                "author": reader.metadata.get("/Author", ""),
                "subject": reader.metadata.get("/Subject", ""),
            }

        text_content = []
        pages_to_extract = min(len(reader.pages), max_pages)

        for page_num in range(pages_to_extract):
            text = reader.pages[page_num].extract_text()
            if text:
                text_content.append(text)

        full_text = "\n\n".join(text_content)

        # Scanned PDFs would need pdf2image + Poppler to rasterize pages
        # for OCR; not wired up, so we report ocr_used=False honestly
        # rather than claiming OCR was attempted.
        if len(full_text.strip()) < 100:
            logger.info("PDF returned minimal extractable text; no OCR fallback configured.")

        return {
            "status": "success",
            "url": url,
            "document_type": "pdf",
            "total_pages": len(reader.pages),
            "pages_extracted": pages_to_extract,
            "metadata": metadata,
            "text": full_text,
            "text_length": len(full_text),
            "ocr_used": False,
            "timestamp": datetime.now().isoformat(),
        }

    async def _analyze_docx(
        self, doc_bytes: bytes, url: str, extract_metadata: bool
    ) -> Dict[str, Any]:
        """Analyze Word document."""
        try:
            from docx import Document
        except ImportError:
            return {
                "status": "partial",
                "document_type": "docx",
                "text": "Word document support requires python-docx library",
                "url": url,
            }

        doc_file = io.BytesIO(doc_bytes)
        document = Document(doc_file)

        text_content = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        full_text = "\n\n".join(text_content)

        metadata = {}
        if extract_metadata:
            core_props = document.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
            }

        return {
            "status": "success",
            "url": url,
            "document_type": "docx",
            "paragraphs": len(document.paragraphs),
            "metadata": metadata,
            "text": full_text,
            "text_length": len(full_text),
            "timestamp": datetime.now().isoformat(),
        }

    async def _analyze_text(self, doc_bytes: bytes, url: str, doc_type: str) -> Dict[str, Any]:
        """Analyze text-based document."""
        try:
            text = doc_bytes.decode("utf-8")
        except UnicodeDecodeError:
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    text = doc_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return {"status": "error", "error": "Unable to decode text document", "url": url}

        return {
            "status": "success",
            "url": url,
            "document_type": doc_type,
            "text": text,
            "text_length": len(text),
            "lines": len(text.split("\n")),
            "timestamp": datetime.now().isoformat(),
        }

    async def _analyze_image(self, image_bytes: bytes, url: str) -> Dict[str, Any]:
        """Analyze image with OCR."""
        if not self.ocr_available:
            return {
                "status": "error",
                "error": "OCR not available. Install pillow and easyocr.",
                "url": url,
            }

        try:
            import easyocr
            import numpy as np
            from PIL import Image

            image = Image.open(io.BytesIO(image_bytes))

            # EasyOCR returns (bbox, text, confidence) tuples from readtext;
            # detail=0 + paragraph=True collapses that to joined strings.
            reader = easyocr.Reader(["en"], gpu=False)
            fragments = reader.readtext(np.array(image), detail=0, paragraph=True)
            text = "\n".join(fragments)

            logger.info("OCR extracted %d characters from image", len(text))

            return {
                "status": "success",
                "url": url,
                "document_type": "image",
                "text": text,
                "text_length": len(text),
                "image_size": f"{image.width}x{image.height}",
                "image_format": image.format,
                "ocr_used": True,
                "timestamp": datetime.now().isoformat(),
            }

        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            return {"status": "error", "error": str(e), "url": url}
